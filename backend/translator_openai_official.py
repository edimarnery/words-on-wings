# -*- coding: utf-8 -*-
"""
Tradutor profissional baseado nas orientações oficiais da OpenAI
Usa Responses API com Structured Outputs para máxima confiabilidade
"""

import os
import sys
import json
import time
import math
import pathlib
import logging
from typing import List, Dict, Optional
from dataclasses import dataclass
from openai import OpenAI
from docx import Document
from tqdm import tqdm

logger = logging.getLogger(__name__)

# Configurações
MODEL = os.getenv("OPENAI_MODEL", "gpt-4.1")
BATCH_TOKEN_BUDGET = int(os.getenv("BATCH_TOKEN_BUDGET", "80000"))
MAX_RETRIES = int(os.getenv("MAX_RETRIES", "6"))
RETRY_BASE_S = float(os.getenv("RETRY_BASE_S", "2.0"))

@dataclass
class TranslationResult:
    success: bool
    translated_segments: int
    processing_time: float
    errors: List[str]
    warnings: List[str]
    checkpoint_path: Optional[str] = None

def estimate_tokens(text: str) -> int:
    """Estimativa conservadora: ~4 chars = 1 token"""
    return max(1, math.ceil(len(text) / 4))

def iter_paragraphs_everywhere(doc: Document):
    """Itera sobre todos os parágrafos do documento (corpo, tabelas, cabeçalhos, rodapés)"""
    # Corpo principal
    for p in doc.paragraphs:
        yield ("body", None, p)
    
    # Tabelas do corpo
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    yield ("table", None, p)
    
    # Cabeçalhos e rodapés por seção
    for sec_idx, sec in enumerate(doc.sections):
        # Cabeçalho
        if hasattr(sec, 'header'):
            for p in sec.header.paragraphs:
                yield ("header", sec_idx, p)
            for t in sec.header.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            yield ("header_table", sec_idx, p)
        
        # Rodapé
        if hasattr(sec, 'footer'):
            for p in sec.footer.paragraphs:
                yield ("footer", sec_idx, p)
            for t in sec.footer.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            yield ("footer_table", sec_idx, p)

def iter_runs_everywhere(doc: Document):
    """Itera sobre todos os runs do documento para preservação fiel de formatação"""
    run_id = 0
    for scope, scope_id, p in iter_paragraphs_everywhere(doc):
        for run in p.runs:
            if run.text and run.text.strip():
                yield (scope, scope_id, p, run, f"r{run_id}")
                run_id += 1

def coletar_runs(doc: Document) -> List[Dict]:
    """Coleta todos os runs com texto para tradução preservando formatação"""
    runs = []
    for scope, scope_id, p, run, run_id in iter_runs_everywhere(doc):
        text = run.text.strip()
        if text:
            runs.append({
                "id": run_id,
                "text": text,
                "scope": scope,
                "scope_id": scope_id
            })
    return runs

def montar_lotes(items: List[Dict], token_budget: int = BATCH_TOKEN_BUDGET) -> List[List[Dict]]:
    """Agrupa items respeitando limite de tokens por lote"""
    lotes, atual, tokens = [], [], 0
    
    for item in items:
        item_tokens = estimate_tokens(item["text"])
        if atual and tokens + item_tokens > token_budget:
            lotes.append(atual)
            atual, tokens = [], 0
        atual.append(item)
        tokens += item_tokens
    
    if atual:
        lotes.append(atual)
    
    return lotes

def pedir_traducao_structured(lote: List[Dict], target_lang: str, source_lang: str = "auto") -> Dict[str, str]:
    """
    Usa Chat Completions API com Structured Outputs - método correto
    """
    client = OpenAI()
    
    # JSON Schema para structured output
    schema = {
        "type": "object",
        "properties": {
            "translations": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "id": {"type": "string"},
                        "translated_text": {"type": "string"}
                    },
                    "required": ["id", "translated_text"],
                    "additionalProperties": False
                }
            }
        },
        "required": ["translations"],
        "additionalProperties": False
    }

    # Prompt otimizado
    prompt = f"""Você é um tradutor profissional especializado. Traduza integralmente cada segmento de texto de {source_lang} para {target_lang}.

INSTRUÇÕES CRÍTICAS:
- NÃO resuma, NÃO omita, NÃO abrevie nenhum conteúdo
- Preserve EXATAMENTE todos os números, datas, siglas e formatação
- Mantenha a mesma estrutura e pontuação
- Traduza PALAVRA POR PALAVRA quando necessário para fidelidade total
- Para termos técnicos, use a tradução padrão mais precisa

Responda APENAS em JSON no formato especificado.

SEGMENTOS PARA TRADUZIR:
{json.dumps([{"id": item["id"], "text": item["text"]} for item in lote], ensure_ascii=False)}"""

    # Retry com backoff exponencial
    for attempt in range(MAX_RETRIES):
        try:
            response = client.chat.completions.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": "Você é um tradutor profissional especializado. Siga exatamente as instruções fornecidas."},
                    {"role": "user", "content": prompt}
                ],
                response_format={
                    "type": "json_schema", 
                    "json_schema": {
                        "name": "translation_result", 
                        "schema": schema
                    }
                },
                temperature=0.1,  # Baixa para consistência
                max_tokens=200000
            )
            
            result_json = json.loads(response.choices[0].message.content)
            return {item["id"]: item["translated_text"] for item in result_json["translations"]}
            
        except Exception as e:
            wait_time = RETRY_BASE_S * (2 ** attempt)
            logger.warning(f"Erro na tentativa {attempt + 1}: {e}")
            
            if attempt == MAX_RETRIES - 1:
                logger.error(f"Falha após {MAX_RETRIES} tentativas: {e}")
                raise
                
            logger.info(f"Aguardando {wait_time}s antes da próxima tentativa...")
            time.sleep(wait_time)

def aplicar_traducoes_runs(doc: Document, traducoes: Dict[str, str]):
    """Aplica traduções preservando formatação run-a-run"""
    for scope, scope_id, p, run, run_id in iter_runs_everywhere(doc):
        if run_id in traducoes:
            run.text = traducoes[run_id]

def garantir_diretorio(path: pathlib.Path):
    """Garante que o diretório existe"""
    path.parent.mkdir(parents=True, exist_ok=True)

def salvar_checkpoint(checkpoint_path: pathlib.Path, traducoes: Dict[str, str]):
    """Salva checkpoint incremental"""
    garantir_diretorio(checkpoint_path)
    with checkpoint_path.open("a", encoding="utf-8") as f:
        f.write(json.dumps(traducoes, ensure_ascii=False) + "\n")

def carregar_checkpoint(checkpoint_path: pathlib.Path) -> Dict[str, str]:
    """Carrega checkpoint existente"""
    traducoes = {}
    if checkpoint_path.exists():
        with checkpoint_path.open("r", encoding="utf-8") as f:
            for linha in f:
                try:
                    lote_traducoes = json.loads(linha.strip())
                    traducoes.update(lote_traducoes)
                except:
                    continue
    return traducoes

def translate_docx_professional(
    input_path: str, 
    output_path: str, 
    source_lang: str, 
    target_lang: str
) -> TranslationResult:
    """
    Tradução profissional de DOCX seguindo orientações oficiais OpenAI
    """
    start_time = time.time()
    errors = []
    warnings = []
    
    try:
        # Paths
        input_path_obj = pathlib.Path(input_path)
        output_path_obj = pathlib.Path(output_path)
        checkpoint_path = pathlib.Path(".checkpoints") / f"{input_path_obj.stem}_runs.jsonl"
        
        # Carregar documento
        logger.info(f"Carregando documento: {input_path}")
        doc = Document(input_path)
        
        # Coletar runs para tradução
        runs = coletar_runs(doc)
        if not runs:
            warnings.append("Nenhum texto encontrado para traduzir")
            doc.save(output_path)
            return TranslationResult(
                success=True,
                translated_segments=0,
                processing_time=time.time() - start_time,
                errors=errors,
                warnings=warnings
            )
        
        logger.info(f"Encontrados {len(runs)} runs de texto para traduzir")
        
        # Carregar checkpoint se existir
        traducoes_existentes = carregar_checkpoint(checkpoint_path)
        runs_pendentes = [r for r in runs if r["id"] not in traducoes_existentes]
        
        if traducoes_existentes:
            logger.info(f"Carregados {len(traducoes_existentes)} runs do checkpoint")
        
        if not runs_pendentes:
            logger.info("Todas as traduções já existem no checkpoint")
            aplicar_traducoes_runs(doc, traducoes_existentes)
            doc.save(output_path)
            return TranslationResult(
                success=True,
                translated_segments=len(runs),
                processing_time=time.time() - start_time,
                errors=errors,
                warnings=warnings,
                checkpoint_path=str(checkpoint_path)
            )
        
        # Processar em lotes
        lotes = montar_lotes(runs_pendentes, BATCH_TOKEN_BUDGET)
        logger.info(f"Processando {len(runs_pendentes)} runs em {len(lotes)} lotes")
        
        traducoes_completas = traducoes_existentes.copy()
        
        for i, lote in enumerate(tqdm(lotes, desc="Traduzindo lotes")):
            try:
                logger.info(f"Traduzindo lote {i+1}/{len(lotes)} ({len(lote)} runs)")
                traducoes_lote = pedir_traducao_structured(lote, target_lang, source_lang)
                traducoes_completas.update(traducoes_lote)
                
                # Salvar checkpoint
                salvar_checkpoint(checkpoint_path, traducoes_lote)
                
            except Exception as e:
                error_msg = f"Erro no lote {i+1}: {e}"
                logger.error(error_msg)
                errors.append(error_msg)
                continue
        
        # Aplicar todas as traduções
        logger.info("Aplicando traduções ao documento...")
        aplicar_traducoes_runs(doc, traducoes_completas)
        
        # Salvar documento final
        garantir_diretorio(output_path_obj)
        doc.save(output_path)
        
        processing_time = time.time() - start_time
        logger.info(f"Tradução concluída em {processing_time:.2f}s")
        
        return TranslationResult(
            success=True,
            translated_segments=len(traducoes_completas),
            processing_time=processing_time,
            errors=errors,
            warnings=warnings,
            checkpoint_path=str(checkpoint_path)
        )
        
    except Exception as e:
        error_msg = f"Erro fatal na tradução: {e}"
        logger.error(error_msg)
        errors.append(error_msg)
        
        return TranslationResult(
            success=False,
            translated_segments=0,
            processing_time=time.time() - start_time,
            errors=errors,
            warnings=warnings
        )

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Uso: python translator_openai_official.py entrada.docx saida.docx pt-BR")
        sys.exit(1)
    
    entrada, saida, target = sys.argv[1], sys.argv[2], sys.argv[3]
    result = translate_docx_professional(entrada, saida, "auto", target)
    
    if result.success:
        print(f"✅ Tradução concluída: {result.translated_segments} segmentos em {result.processing_time:.2f}s")
    else:
        print(f"❌ Falha na tradução: {result.errors}")
        sys.exit(1)